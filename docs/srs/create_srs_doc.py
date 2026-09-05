from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import glob

doc = Document()

# ตั้ง font เริ่มต้น
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# โหลดไฟล์ SRS ทั้งหมด
srs_folder = 'docs/srs'
files = sorted(glob.glob(os.path.join(srs_folder, '*.md')))

for file_path in files:
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # หาชื่อระบบจากหัวข้อแรก
    title = os.path.basename(file_path).replace('-srs.md', '').replace('-', ' ').title()
    
    # เพิ่มหัวข้อหลัก
    doc.add_heading(title, level=0)
    doc.add_page_break()
    
    # แยกย่อหน้า
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('```'):
            continue  # ข้าม code block markers
        elif line.startswith('|'):
            # ตาราง
            doc.add_paragraph(line)
        elif line.strip() == '':
            continue
        else:
            doc.add_paragraph(line.strip())
    
    doc.add_page_break()

doc.save('SRS_Kinyrak_Ingredients.docx')
print("บันทึกเรียบร้อยแล้ว!")