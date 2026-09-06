import os
import glob
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

srs_folder = 'docs/srs'
files = sorted(glob.glob(os.path.join(srs_folder, '*.md')))

for file_path in files:
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    file_name = os.path.basename(file_path)
    title = file_name.replace('-srs.md', '').replace('-', ' ').title()

    doc.add_heading(title, level=0)

    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()

        if stripped.startswith('```'):
            continue
        elif stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            heading_text = stripped.lstrip('#').strip()
            if level == 1:
                doc.add_heading(heading_text, level=1)
            elif level == 2:
                doc.add_heading(heading_text, level=2)
            elif level == 3:
                doc.add_heading(heading_text, level=3)
            else:
                doc.add_paragraph(heading_text)
        elif stripped.startswith('|'):
            if set(stripped.replace('|', '').replace('-', '').replace(' ', '').strip()) == set():
                continue
            parts = [p.strip() for p in stripped.split('|') if p.strip()]
            if parts:
                p = doc.add_paragraph(' | '.join(parts))
                p.paragraph_format.space_after = Pt(2)
        elif stripped == '':
            continue
        elif stripped.startswith('- '):
            text = stripped[2:]
            p = doc.add_paragraph(text, style='List Bullet')
        elif stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph(stripped.strip('*'))
            for run in p.runs:
                run.bold = True
        else:
            p = doc.add_paragraph(stripped)

    doc.add_page_break()

doc.save('SRS_Kinyrak_Ingredients.docx')
print("บันทึกเรียบร้อยแล้ว: SRS_Kinyrak_Ingredients.docx")
